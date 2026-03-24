from io import BytesIO
import re

import pdfplumber
from docx import Document


def _clean_text_blocks(blocks: list[str]) -> str:
    cleaned_blocks = [block.strip() for block in blocks if block and block.strip()]
    text = "\n\n".join(cleaned_blocks)
    text = re.sub(r"(?<=\w)-\n(?=\w)", "", text)
    text = re.sub(r"(?<!\n)\n(?!\n)", " ", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_pdf_text_and_tables(file_bytes: bytes) -> tuple[str, list[str], bool]:
    page_chunks: list[str] = []
    table_chunks: list[str] = []
    likely_scanned = False

    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        page_count = max(1, len(pdf.pages))
        low_text_pages = 0

        for page in pdf.pages:
            plain = (page.extract_text(layout=False) or "").strip()
            layout = (page.extract_text(layout=True) or "").strip()
            best_text = layout if len(layout) > len(plain) else plain

            if len(best_text) < 40:
                low_text_pages += 1

            if best_text:
                page_chunks.append(best_text)

            tables = page.extract_tables() or []
            for table in tables:
                if not table:
                    continue
                rows = []
                for row in table:
                    cells = [str(cell).strip() for cell in row if cell and str(cell).strip()]
                    if cells:
                        rows.append(" | ".join(cells))
                if rows:
                    table_chunks.append("\n".join(rows))

        likely_scanned = low_text_pages / page_count >= 0.6

    combined = page_chunks + (["TABLE DATA:\n" + "\n\n".join(table_chunks)] if table_chunks else [])
    return _clean_text_blocks(combined), table_chunks, likely_scanned


def parse_pdf_with_diagnostics(file_bytes: bytes) -> tuple[str, dict]:
    diagnostics = {
        "parser": "pdfplumber",
        "warnings": [],
        "tables_detected": 0,
        "scanned_pdf_likely": False,
    }

    try:
        text, tables, likely_scanned = _extract_pdf_text_and_tables(file_bytes)
        diagnostics["tables_detected"] = len(tables)
        diagnostics["scanned_pdf_likely"] = likely_scanned

        if likely_scanned:
            diagnostics["warnings"].append(
                "PDF appears scanned/image-based. OCR preprocessing is recommended for best accuracy."
            )
        if tables:
            diagnostics["warnings"].append(
                "Table data was detected and flattened; verify important rows were captured correctly."
            )
        if len(text) < 250:
            diagnostics["warnings"].append(
                "Very little text was extracted; formatting quality may affect scoring precision."
            )

        if not text:
            raise ValueError("No readable text found in the uploaded PDF.")
        return text, diagnostics
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Failed to parse PDF file: {exc}") from exc


def parse_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file represented as bytes."""
    text, _diagnostics = parse_pdf_with_diagnostics(file_bytes)
    return text


def parse_docx_with_diagnostics(file_bytes: bytes) -> tuple[str, dict]:
    diagnostics = {
        "parser": "python-docx",
        "warnings": [],
        "tables_detected": 0,
        "scanned_pdf_likely": False,
    }

    try:
        document = Document(BytesIO(file_bytes))
        paragraph_text = [paragraph.text for paragraph in document.paragraphs]

        table_chunks = []
        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if cells:
                    rows.append(" | ".join(cells))
            if rows:
                table_chunks.append("\n".join(rows))

        diagnostics["tables_detected"] = len(table_chunks)
        if table_chunks:
            diagnostics["warnings"].append(
                "DOCX table content was flattened; review alignment-sensitive entries."
            )

        text = _clean_text_blocks(paragraph_text + table_chunks)
        if not text:
            raise ValueError("No readable text found in the uploaded DOCX.")
        return text, diagnostics
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError(f"Failed to parse DOCX file: {exc}") from exc


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file represented as bytes."""
    text, _diagnostics = parse_docx_with_diagnostics(file_bytes)
    return text
