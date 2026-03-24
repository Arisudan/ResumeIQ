from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt


def export_to_docx(optimized_text: str) -> bytes:
    document = Document()

    # Set default font for normal paragraphs.
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    # Configure heading style.
    heading_style = document.styles["Heading 1"]
    heading_style.font.name = "Calibri"
    heading_style.font.size = Pt(13)
    heading_style.font.bold = True

    # Set 1-inch margins on all sides.
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    for raw_line in optimized_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        is_heading = line.isupper() or line.endswith(":")
        is_bullet = line.startswith(("•", "-", "*"))

        if is_heading:
            document.add_paragraph(line.rstrip(":"), style="Heading 1")
        elif is_bullet:
            bullet_text = line.lstrip("•-* ").strip()
            paragraph = document.add_paragraph(bullet_text, style="List Bullet")
            paragraph.paragraph_format.left_indent = Inches(0.25)
        else:
            document.add_paragraph(line, style="Normal")

    output = BytesIO()
    document.save(output)
    output.seek(0)
    return output.read()
